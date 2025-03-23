import os
import fitz  # PyMuPDF for PDF text and image extraction
import pdfplumber  # For PDF text and table extraction
from docx import Document  # For DOCX text and table extraction
import pandas as pd  # For CSV processing
import asyncio  # For asynchronous processing
import aiofiles  # For asynchronous file I/O
import logging
from datetime import datetime
import json

# Configure logging
logging.basicConfig(
    filename="processing.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

UPLOAD_FOLDER = "Data_Folder"
OUTPUT_FOLDER = "Processed_Data"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)


async def process_pdf(pdf_path):
    """
    Asynchronously process a PDF file to extract text, tables, and images.
    """
    text_data, table_data, image_data = [], [], []
    try:
        # Extract text and tables using pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):

                text = page.extract_text()
                if text:
                    text_data.append({"page": page_num, "text": text})

                tables = page.extract_tables()
                for table_num, table in enumerate(tables, start=1):
                    table_data.append({"page": page_num, "table_num": table_num, "table": table})

        pdf_doc = fitz.open(pdf_path)
        for page_num in range(len(pdf_doc)):
            for img_index, img in enumerate(pdf_doc[page_num].get_images(full=True)):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_filename = f"{os.path.basename(pdf_path)}_page{page_num + 1}_img{img_index + 1}.png"
                img_path = os.path.join(OUTPUT_FOLDER, img_filename)

                # Save image to disk asynchronously
                async with aiofiles.open(img_path, "wb") as img_file:
                    await img_file.write(image_bytes)

                image_data.append({"page": page_num + 1, "image_path": img_path})

        logging.info(f"Successfully processed PDF: {pdf_path}")
        return {"file": pdf_path, "text": text_data, "tables": table_data, "images": image_data}

    except Exception as e:
        logging.error(f"Error processing PDF {pdf_path}: {e}")
        return {"file": pdf_path, "error": str(e)}


async def process_docx(docx_path):
    """
    Asynchronously process a DOCX file to extract text, tables, and images.
    """
    text_data, table_data, image_data = [], [], []
    try:
        doc = Document(docx_path)

        # Extract text
        for para in doc.paragraphs:
            text_data.append(para.text)

        # Extract tables
        for table_num, table in enumerate(doc.tables, start=1):
            table_content = []
            for row in table.rows:
                table_content.append([cell.text for cell in row.cells])
            table_data.append({"table_num": table_num, "table": table_content})

        # Extract images
        for rel in doc.part.rels:
            if "image" in doc.part.rels[rel].target_ref:
                img_data = doc.part.rels[rel].target_part.blob
                img_filename = f"{os.path.basename(docx_path)}_img.png"
                img_path = os.path.join(OUTPUT_FOLDER, img_filename)

                # Save image to disk asynchronously
                async with aiofiles.open(img_path, "wb") as img_file:
                    await img_file.write(img_data)

                image_data.append({"image_path": img_path})

        logging.info(f"Successfully processed DOCX: {docx_path}")
        return {"file": docx_path, "text": text_data, "tables": table_data, "images": image_data}

    except Exception as e:
        logging.error(f"Error processing DOCX {docx_path}: {e}")
        return {"file": docx_path, "error": str(e)}


async def process_csv(csv_path):
    """
    Asynchronously process a CSV file to extract tables.
    """
    try:
        # Read CSV in chunks to handle large files
        chunks = pd.read_csv(csv_path, chunksize=10000) 
        table_data = []
        for chunk_num, chunk in enumerate(chunks, start=1):
            table_data.extend(chunk.to_dict(orient="records"))

        logging.info(f"Successfully processed CSV: {csv_path}")
        return {"file": csv_path, "tables": table_data}

    except Exception as e:
        logging.error(f"Error processing CSV {csv_path}: {e}")
        return {"file": csv_path, "error": str(e)}


async def process_all_files(folder_path):
    """
    Asynchronously process all files in the given folder.
    """
    results = []
    files = os.listdir(folder_path)

    # Create a list of tasks for asynchronous processing
    tasks = []
    for file in files:
        file_path = os.path.join(folder_path, file)
        if file.endswith(".pdf"):
            tasks.append(process_pdf(file_path))
        elif file.endswith(".docx"):
            tasks.append(process_docx(file_path))
        elif file.endswith(".csv"):
            tasks.append(process_csv(file_path))

    # Run all tasks concurrently
    results = await asyncio.gather(*tasks)

    # Save extracted data as JSON
    output_file = os.path.join(OUTPUT_FOLDER, "processed_data.json")
    async with aiofiles.open(output_file, "w") as json_file:
        await json_file.write(json.dumps(results, indent=4))

    logging.info(f"Processing complete! Data saved in {output_file}")
    return results

if __name__ == "__main__":
    asyncio.run(process_all_files(UPLOAD_FOLDER))
    print("Processing complete! Data saved in processed_data.json")